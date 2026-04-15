"""
Generate Word (.docx) and PDF for Now Assist vs Custom-Built comparison
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import PageBreak

# ─────────────────────────────────────────────
# DATA
# ─────────────────────────────────────────────

TITLE = "Tool Comparison"
SUBTITLE = "ServiceNow Now Assist vs Custom-Built AI Assistant"
DATE = "April 2026 | LTIMindtree Internal"

SOURCES = [
    ("ServiceNow Now Assist — Official Product Page", "https://www.servicenow.com/platform/now-assist.html"),
    ("Gartner Peer Insights — ServiceNow Now Assist Reviews & Ratings 2026", "https://www.gartner.com/reviews/product/servicenow-now-assist"),
    ("Plat4mation — What is ServiceNow Now Assist and why does it matter?", "https://plat4mation.com/blog/what-is-servicenow-now-assist-and-why-does-it-matter/"),
    ("eesel AI — 6 Best AI tools for ServiceNow in 2026: Complete comparison", "https://www.eesel.ai/blog/best-ai-for-servicenow"),
    ("Crossfuze — Now Assist FAQs", "https://www.crossfuze.com/post/now-assist-faqs"),
    ("Redress Compliance — ServiceNow Now Assist Pricing and the ROI Reality", "https://redresscompliance.com/servicenow-now-assist-ai-strategy-white-paper.html"),
    ("servistio — AI in ServiceNow: Now Assist GenAI vs AI Agents", "https://servistio.com/en/blog/ai-in-servicenow-what-is-the-difference-between-now-assist-genai-and-ai-agents/"),
    ("ServiceNow Community — Now Assist FAQs", "https://www.servicenow.com/community/now-assist-articles/now-assist-faqs/ta-p/2685122"),
]

overview_headers = ["", "Now Assist", "Custom-Built AI Assistant"]
overview_rows = [
    ["What it is", "GenAI layer built natively into ServiceNow platform", "AI assistant built on BlueVerse Foundry, LangChain, or direct LLM APIs"],
    ["Vendor", "ServiceNow (NowLLM + Azure OpenAI / Claude / Gemini)", "Self-built, LTIMindtree-owned"],
    ["Deployment", "SaaS — lives inside ServiceNow only", "Flexible — cloud, on-prem, hybrid"],
    ["Time to deploy", "Weeks (configuration-heavy)", "Varies — days to months depending on complexity"],
]

feature_headers = ["Capability", "Now Assist", "Custom-Built"]
feature_rows = [
    ["Incident / Case Summarization", "✅ Out-of-the-box", "✅ Buildable"],
    ["Chat & Virtual Agent", "✅ Native", "✅ Buildable"],
    ["Email / Content Generation", "✅ Native", "✅ Buildable"],
    ["RAG over company documents", "⚠️ Limited to ServiceNow data only", "✅ Any data source (SharePoint, Confluence, PDFs, DBs)"],
    ["Multi-system knowledge", "❌ ServiceNow-only", "✅ Any system via connectors"],
    ["Custom workflows", "⚠️ Within ServiceNow flows only", "✅ Fully flexible"],
    ["Agent / Agentic workflows", "✅ Available (2025+)", "✅ Fully buildable"],
    ["Model choice", "⚠️ NowLLM, Claude, Azure OpenAI, Gemini (limited control)", "✅ Any LLM — full control"],
    ["Fine-tuning / custom model", "❌ Not supported", "✅ Fully supported"],
    ["Guardrails & governance", "⚠️ Basic, platform-enforced", "✅ Custom — full control"],
    ["Multi-language support", "⚠️ English only (verified)", "✅ Configurable per language"],
    ["Integration with external tools", "⚠️ Limited to ServiceNow ecosystem", "✅ Any API, MCP, tool"],
    ["Analytics & evaluation", "⚠️ Limited native testing tools", "✅ Custom evaluation (e.g., BlueVerse Evaluator)"],
    ["ServiceNow-specific tasks", "✅ Best-in-class", "⚠️ Requires custom integration"],
]

data_headers = ["", "Now Assist", "Custom-Built"]
data_rows = [
    ["ServiceNow records", "✅ Native, deep access", "⚠️ Requires API integration"],
    ["SharePoint / OneDrive", "❌ Not natively", "✅ Via connectors"],
    ["Confluence / Jira", "❌ Not natively", "✅ Via connectors"],
    ["Internal PDFs / Policy docs", "⚠️ Only if uploaded to ServiceNow", "✅ Direct RAG ingestion"],
    ["Custom databases", "❌", "✅"],
    ["Real-time external data", "❌", "✅ Via tools/APIs"],
]

cost_headers = ["", "Now Assist", "Custom-Built"]
cost_rows = [
    ["Pricing model", "Quote-based, bundled into enterprise contracts", "LLM API costs + build/maintain costs"],
    ["Transparency", "❌ Opaque pricing, often expensive add-on", "✅ Pay-per-use, predictable"],
    ["Lock-in risk", "❌ High — all config locked to ServiceNow", "✅ Low — portable"],
    ["Incremental cost per feature", "❌ High — each Now Assist SKU priced separately", "✅ Marginal — extend existing build"],
]

when_headers = ["Scenario", "Recommended"]
when_rows = [
    ["Client is 100% on ServiceNow, no external data needed", "Now Assist"],
    ["Client needs AI across multiple platforms (ServiceNow + SharePoint + Confluence)", "Custom-Built"],
    ["Client wants to own the AI, brand it, and extend it freely", "Custom-Built"],
    ["Quick win for ITSM summarization with no custom needs", "Now Assist"],
    ["Client has sensitive data that cannot go to ServiceNow's cloud", "Custom-Built"],
    ["Client wants to evaluate and test AI quality rigorously", "Custom-Built"],
    ["Client wants full RAG over internal policy/HR/compliance docs", "Custom-Built"],
    ["Already paying for ServiceNow enterprise, wants low-friction AI", "Now Assist"],
]

verdict_headers = ["Dimension", "Winner"]
verdict_rows = [
    ["Speed to deploy (ITSM-specific)", "Now Assist"],
    ["Flexibility & extensibility", "Custom-Built"],
    ["Data coverage", "Custom-Built"],
    ["Cost transparency", "Custom-Built"],
    ["ServiceNow-specific tasks", "Now Assist"],
    ["Model control", "Custom-Built"],
    ["Evaluation & governance", "Custom-Built"],
    ["Vendor lock-in risk", "Custom-Built"],
]

now_assist_limits = [
    "Platform lock-in — works only inside ServiceNow. If client moves platforms, AI investment is lost.",
    "10–20% real resolution rate — ServiceNow markets 40–60% but real-world deployments show 10–20%.",
    "Hallucination risk — limits use to low-risk, well-defined workflows.",
    "Complex setup — slow time-to-value despite being 'out-of-the-box'.",
    "No external knowledge — cannot access Confluence, Google Docs, SharePoint, or non-ServiceNow sources natively.",
    "Limited model control — customers cannot freely choose or switch LLMs.",
    "English-only verified — multilingual support is experimental and unverified.",
]

custom_limits = [
    "Build effort — requires upfront development investment.",
    "ServiceNow deep integration — accessing ServiceNow records requires API work.",
    "Maintenance — team must own updates, model upgrades, and monitoring.",
    "Governance burden — guardrails, audit, and compliance must be built from scratch.",
]

bottom_line = (
    "Now Assist wins for pure ServiceNow ITSM use cases where speed and native integration matter. "
    "Custom-Built wins everywhere else — multi-system knowledge, branded experiences, full control over "
    "prompts/models/evaluation, and any scenario where the client's data lives outside ServiceNow."
)

# ─────────────────────────────────────────────
# WORD DOCUMENT
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, '1B3A6B')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(h)
        run.text = h
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = 'F0F4FA' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = val
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run(val)
            run.text = val
            run.font.size = Pt(8.5)
            if c_idx == 0:
                run.bold = True

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])

    return table

def build_word():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run(SUBTITLE)
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = date_p.add_run(DATE)
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    def section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    # Overview
    section_heading("Overview")
    add_table(doc, overview_headers, overview_rows, col_widths=[1.5, 2.8, 2.8])
    doc.add_paragraph()

    # Feature Comparison
    section_heading("Feature-by-Feature Comparison")
    add_table(doc, feature_headers, feature_rows, col_widths=[2.2, 2.4, 2.5])
    doc.add_paragraph()

    # Data & Knowledge
    section_heading("Data & Knowledge Access")
    add_table(doc, data_headers, data_rows, col_widths=[2.0, 2.5, 2.5])
    doc.add_paragraph()

    # Cost
    section_heading("Cost & Licensing")
    add_table(doc, cost_headers, cost_rows, col_widths=[2.0, 2.5, 2.5])
    doc.add_paragraph()

    # Limitations — Now Assist
    section_heading("Limitations of Now Assist")
    for item in now_assist_limits:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(9.5)

    doc.add_paragraph()

    # Limitations — Custom
    section_heading("Limitations of Custom-Built")
    for item in custom_limits:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(9.5)

    doc.add_paragraph()

    # When to choose
    section_heading("When to Choose Which")
    add_table(doc, when_headers, when_rows, col_widths=[4.0, 2.0])
    doc.add_paragraph()

    # Verdict
    section_heading("Summary Verdict")
    add_table(doc, verdict_headers, verdict_rows, col_widths=[3.0, 3.0])
    doc.add_paragraph()

    # Bottom line
    p = doc.add_paragraph()
    run = p.add_run("Bottom Line: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    run2 = p.add_run(bottom_line)
    run2.font.size = Pt(10)

    # Sources
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Sources & References")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    for i, (label, url) in enumerate(SOURCES, start=1):
        p = doc.add_paragraph(style='List Number')
        run_label = p.add_run(f"{label} — ")
        run_label.font.size = Pt(9)
        run_url = p.add_run(url)
        run_url.font.size = Pt(9)
        run_url.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

    doc.save("NOW_ASSIST_VS_CUSTOM_BUILT_COMPARISON.docx")
    print("Word document saved.")

# ─────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────

DARK_BLUE = colors.HexColor('#1B3A6B')
MID_BLUE  = colors.HexColor('#2E86C1')
LIGHT_BG  = colors.HexColor('#F0F4FA')
WHITE     = colors.white
GREY_TEXT = colors.HexColor('#555555')
GREEN     = colors.HexColor('#1A7A4A')
RED       = colors.HexColor('#C0392B')
AMBER     = colors.HexColor('#D68910')

def build_pdf():
    doc = SimpleDocTemplate(
        "NOW_ASSIST_VS_CUSTOM_BUILT_COMPARISON.pdf",
        pagesize=A4,
        topMargin=1.8*cm, bottomMargin=1.8*cm,
        leftMargin=2*cm, rightMargin=2*cm,
        title="Now Assist vs Custom-Built AI Assistant"
    )

    styles = getSampleStyleSheet()

    style_title = ParagraphStyle('DocTitle', fontSize=22, textColor=DARK_BLUE,
                                  alignment=TA_CENTER, spaceAfter=6, fontName='Helvetica-Bold')
    style_subtitle = ParagraphStyle('DocSubtitle', fontSize=13, textColor=MID_BLUE,
                                     alignment=TA_CENTER, spaceAfter=4, fontName='Helvetica-Bold')
    style_date = ParagraphStyle('DocDate', fontSize=8, textColor=GREY_TEXT,
                                 alignment=TA_CENTER, spaceAfter=14)
    style_section = ParagraphStyle('Section', fontSize=11, textColor=DARK_BLUE,
                                    spaceBefore=14, spaceAfter=5, fontName='Helvetica-Bold')
    style_body = ParagraphStyle('Body', fontSize=8.5, textColor=colors.black,
                                 spaceAfter=3, leading=13)
    style_bullet = ParagraphStyle('Bullet', fontSize=8.5, textColor=colors.black,
                                   spaceAfter=3, leftIndent=14, leading=13)
    style_bottom = ParagraphStyle('Bottom', fontSize=9.5, textColor=DARK_BLUE,
                                   spaceBefore=8, leading=14, borderPad=6,
                                   backColor=LIGHT_BG, borderWidth=1,
                                   borderColor=MID_BLUE)

    def make_table(headers, rows, col_widths):
        data = [headers] + rows
        tbl = Table(data, colWidths=col_widths, repeatRows=1)

        style_cmds = [
            # Header
            ('BACKGROUND', (0,0), (-1,0), DARK_BLUE),
            ('TEXTCOLOR', (0,0), (-1,0), WHITE),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 8),
            ('ALIGN', (0,0), (-1,0), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('FONTSIZE', (0,1), (-1,-1), 8),
            ('FONTNAME', (0,1), (0,-1), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 0.4, colors.HexColor('#CCCCCC')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [LIGHT_BG, WHITE]),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('LEFTPADDING', (0,0), (-1,-1), 5),
            ('RIGHTPADDING', (0,0), (-1,-1), 5),
        ]

        # Colour Now Assist / Custom-Built verdict cells
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, val in enumerate(row):
                if val in ("Now Assist", "Custom-Built"):
                    bg = MID_BLUE if val == "Now Assist" else GREEN
                    style_cmds.append(('BACKGROUND', (c_idx, r_idx), (c_idx, r_idx), bg))
                    style_cmds.append(('TEXTCOLOR', (c_idx, r_idx), (c_idx, r_idx), WHITE))
                    style_cmds.append(('FONTNAME', (c_idx, r_idx), (c_idx, r_idx), 'Helvetica-Bold'))

        tbl.setStyle(TableStyle(style_cmds))
        return tbl

    story = []

    # Title block
    story.append(Paragraph(TITLE, style_title))
    story.append(Paragraph(SUBTITLE, style_subtitle))
    story.append(Paragraph(DATE, style_date))
    story.append(HRFlowable(width="100%", thickness=1.5, color=DARK_BLUE))
    story.append(Spacer(1, 10))

    PAGE_W = A4[0] - 4*cm  # usable width

    # Overview
    story.append(Paragraph("Overview", style_section))
    story.append(make_table(overview_headers, overview_rows,
                             [3.5*cm, 7.5*cm, 7.5*cm]))
    story.append(Spacer(1, 10))

    # Feature comparison
    story.append(Paragraph("Feature-by-Feature Comparison", style_section))
    story.append(make_table(feature_headers, feature_rows,
                             [4.8*cm, 6.8*cm, 6.9*cm]))
    story.append(Spacer(1, 10))

    # Data
    story.append(Paragraph("Data & Knowledge Access", style_section))
    story.append(make_table(data_headers, data_rows,
                             [4.5*cm, 7*cm, 7*cm]))
    story.append(Spacer(1, 10))

    # Cost
    story.append(Paragraph("Cost & Licensing", style_section))
    story.append(make_table(cost_headers, cost_rows,
                             [4.5*cm, 7*cm, 7*cm]))
    story.append(Spacer(1, 10))

    # Limitations
    story.append(Paragraph("Limitations of Now Assist", style_section))
    for item in now_assist_limits:
        story.append(Paragraph(f"• {item}", style_bullet))

    story.append(Spacer(1, 6))
    story.append(Paragraph("Limitations of Custom-Built", style_section))
    for item in custom_limits:
        story.append(Paragraph(f"• {item}", style_bullet))

    story.append(Spacer(1, 10))

    # When to choose
    story.append(Paragraph("When to Choose Which", style_section))
    story.append(make_table(when_headers, when_rows,
                             [11.5*cm, 7*cm]))
    story.append(Spacer(1, 10))

    # Verdict
    story.append(Paragraph("Summary Verdict", style_section))
    story.append(make_table(verdict_headers, verdict_rows,
                             [9.25*cm, 9.25*cm]))
    story.append(Spacer(1, 12))

    # Bottom line
    story.append(HRFlowable(width="100%", thickness=1, color=MID_BLUE))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        f"<b>Bottom Line:</b> {bottom_line}",
        style_bottom
    ))

    # Sources
    story.append(Spacer(1, 14))
    story.append(HRFlowable(width="100%", thickness=1, color=DARK_BLUE))
    story.append(Spacer(1, 6))
    story.append(Paragraph("Sources & References", style_section))

    style_source = ParagraphStyle('Source', fontSize=8, textColor=colors.black,
                                   spaceAfter=3, leading=12, leftIndent=10)
    style_source_url = ParagraphStyle('SourceURL', fontSize=7.5, textColor=MID_BLUE,
                                       spaceAfter=5, leading=11, leftIndent=20)

    for i, (label, url) in enumerate(SOURCES, start=1):
        story.append(Paragraph(f"{i}. <b>{label}</b>", style_source))
        story.append(Paragraph(url, style_source_url))

    doc.build(story)
    print("PDF saved.")

# ─────────────────────────────────────────────
# RUN
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import os
    os.chdir(r"C:\Users\Admin\Desktop\blueverse")
    build_word()
    build_pdf()
    print("Done. Both files saved to Desktop/blueverse/")
